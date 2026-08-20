"""
Fix -- Antagonist Rotation Rule violations, Kelly Mulroy 5-Day Training Plan, Day 2
(Wednesday, 70% Green). Found by the 8/15/2026 rotating daily-audit check-in.

Kelly Mulroy has no build script in this repo -- clients/kelly_mulroy/Kelly_Mulroy_5Day_
Training_Plan.docx is the hand-delivered original reference document the whole
icons_template.js engine was XML-audited against (see CLAUDE.md's "Canonical reference").
There is nothing to regenerate it from, so this script edits the delivered .docx directly
and surgically via python-docx, rather than hand-writing the output file -- it locates the
two violating exercise tables and swaps the full row content (all 7 columns) of exactly the
two rows involved, leaving every other row/cell in the document completely untouched.

Violations found:
  - Block A "Primary Press Strength": Dumbbell Chest Press (Flat) -> Incline Dumbbell Press
    -> Dumbbell Overhead Press -- 3 consecutive push/press exercises.
  - Block B "Primary Pull Strength": Dumbbell Single-Arm Row (Left First) -> Lat Pulldown
    (Cable) -> Face Pull (Cable/Band) -- 3 consecutive pull exercises.

Fix: swap Block A's 3rd exercise (Dumbbell Overhead Press, a push exercise) with Block B's
1st exercise (Dumbbell Single-Arm Row (Left First), a pull exercise) -- a single cross-block
row swap that resolves BOTH flagged violations in one move, with no new same-pattern run
introduced at the block boundary:

    Block A becomes: Chest Press (push) -> Incline DB Press (push) -> Single-Arm Row (pull)
    Block B becomes: Overhead Press (push) -> Lat Pulldown (pull) -> Face Pull (pull)
    Full Day-2 sequence across the A/B boundary: Push, Push, Pull | Push, Pull, Pull
    -- no 3-in-a-row anywhere, including across the block boundary.

Every other same-block alternative considered (e.g. swapping OHP with Face Pull directly,
matching the "3rd exercise" position in both blocks) was checked and rejected because it
recreates a 3-consecutive-pull run right at the A/B boundary (Face Pull, Row, Lat Pulldown).
The 1st<->3rd cross swap used here is the only single-swap fix that clears the violation in
both blocks without creating a new one.

No exercise is dropped; no set/rep/load/tempo/rest/cue value is changed -- the two exercises
(and their own full row of prescription data) simply trade table positions.
"""
import docx

PATH = "clients/kelly_mulroy/Kelly_Mulroy_5Day_Training_Plan.docx"

d = docx.Document(PATH)

# Confirmed via prior structural read: table index 13 = Day 2 Block A "Primary Press
# Strength" exercise body table; table index 15 = Day 2 Block B "Primary Pull Strength"
# exercise body table.
block_a = d.tables[13]
block_b = d.tables[15]

a_row = block_a.rows[2]  # Dumbbell Overhead Press (violating 3rd push exercise)
b_row = block_b.rows[0]  # Dumbbell Single-Arm Row (Left First) (nearby pull exercise)

a_name = a_row.cells[0].paragraphs[0].runs[0].text
b_name = b_row.cells[0].paragraphs[0].runs[0].text
assert a_name == "Dumbbell Overhead Press", a_name
assert b_name == "Dumbbell Single-Arm Row (Left First)", b_name

NUM_COLS = 7
for col in range(NUM_COLS):
    a_cell = a_row.cells[col]
    b_cell = b_row.cells[col]
    a_run = a_cell.paragraphs[0].runs[0]
    b_run = b_cell.paragraphs[0].runs[0]
    a_text, b_text = a_run.text, b_run.text
    a_run.text = b_text
    b_run.text = a_text

d.save(PATH)

# Verify
d2 = docx.Document(PATH)
new_a_row = [c.text for c in d2.tables[13].rows[2].cells]
new_b_row = [c.text for c in d2.tables[15].rows[0].cells]
print("Block A row 3 (was OHP) is now:", new_a_row[0])
print("Block B row 1 (was Row) is now:", new_b_row[0])
assert new_a_row[0] == "Dumbbell Single-Arm Row (Left First)"
assert new_b_row[0] == "Dumbbell Overhead Press"
print("OK -- swap verified.")
