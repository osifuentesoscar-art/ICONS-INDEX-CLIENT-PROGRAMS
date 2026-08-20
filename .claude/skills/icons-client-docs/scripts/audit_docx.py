#!/usr/bin/env python3
"""Structural audit for ICONS .docx deliverables.

LibreOffice headless conversion is broken in this environment, so a rendered
visual check isn't available. This extracts every paragraph and table cell and
checks the things that actually break: placeholder leakage from an incomplete
data object, missing auto-fired clinical callouts, and sections a comparable
document would have.

    python3 audit_docx.py <file.docx> [--type plan|assessment]
                          [--compare <sibling.docx>] [--expect TERM ...]

Exit code is 1 if any hard failure is found, so it can gate a delivery step.
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.exit("python-docx not installed.  pip install python-docx")

# Template artifacts that are unambiguous: JS renders these literally when a
# data field is missing, and they do not occur in written English.  Matched
# case-sensitively on purpose -- see SOFT_PATTERNS for why.
HARD_PATTERNS = [
    (r"\bundefined\b", "undefined"),
    (r"\bNaN\b", "NaN"),
    (r"\[object Object\]", "[object Object]"),
    (r"\$\{", "unevaluated ${} template literal"),
    (r"\{\{", "unfilled {{ }} placeholder"),
]

# Markers that are real when written as a placeholder (conventionally uppercase)
# but ordinary vocabulary in lowercase prose -- an ICONS Note legitimately says
# "not a placeholder", and flagging that trains the reader to ignore the audit.
# Case-sensitive uppercase matching separates the two cleanly.
SOFT_PATTERNS = [
    (r"PLACEHOLDER", "PLACEHOLDER marker"),
    (r"\bTODO\b", "TODO"),
    (r"\bTBD\b", "TBD"),
    (r"\bX{2,}\b", "XX placeholder"),
    (r"\bnull\b", "null"),
    (r"\bLorem ipsum\b", "lorem ipsum"),
]

# Callouts the engine inserts from client fields, plus the section headings a
# well-formed plan carries.  Absence is a signal to check the data object, not
# proof of a bug -- reported as warnings.
DOC_PROFILES = {
    "plan": {
        "callouts": {
            "protein bar (ALST At-Risk)": ["protein", "g/day"],
            "pelvic floor callout": ["pelvic floor"],
            "warm-up": ["warm-up"],
            "cool-down": ["cool-down"],
            "ICONS note": ["icons note"],
            "progression / RIR": ["rir"],
        },
        "sections": {
            "baselines table": ["baseline"],
            "weekly summary": ["progression target", "key lift"],
            "milestones": ["4-week", "8-week", "4 week", "8 week"],
            "rescan note": ["rescan"],
        },
    },
    "assessment": {
        "callouts": {
            "biological age disclaimer": ["biological age"],
            "segmental distribution": ["lean mass", "left arm", "left leg"],
            "strength assessment table": ["% bw", "level"],
            "footnotes appendix": ["methodology", "how to read"],
        },
        "sections": {
            "styku stat grid": ["body fat", "shape score"],
            "benefit cards": ["aesthetics", "health"],
            "next steps": ["next step", "reassessment"],
        },
    },
}


def detect_type(path, text):
    """Assessment reports and training plans expect different sections, so
    checking one against the other's profile produces pure noise."""
    name = Path(path).name.lower()
    if "assessment" in name or "performance_assessment" in name:
        return "assessment"
    if "improvement" in name:
        return "assessment"
    if "training_plan" in name or "plan" in name:
        return "plan"
    return "assessment" if "biological age" in text else "plan"


def extract(path):
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs]
    cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cells.append(cell.text)
    return doc, paras, cells


def scan(chunks, patterns):
    hits = []
    for chunk in chunks:
        for pattern, label in patterns:
            m = re.search(pattern, chunk)
            if m:
                start = max(0, m.start() - 60)
                hits.append((label, chunk.strip()[start:m.end() + 60]))
    return hits


def audit(path, compare=None, expect=(), doctype=None):
    doc, paras, cells = extract(path)
    chunks = paras + cells
    text = "\n".join(chunks).lower()

    failures, warnings = [], []
    doctype = doctype or detect_type(path, text)
    profile = DOC_PROFILES[doctype]

    print(f"\n=== {Path(path).name} ===")
    print(f"type: {doctype}   paragraphs: {len(doc.paragraphs)}   "
          f"tables: {len(doc.tables)}   table cells: {len(cells)}")

    # --- hard failures -----------------------------------------------------
    hard = scan(chunks, HARD_PATTERNS)
    if hard:
        print("\n  TEMPLATE LEAKAGE")
        for label, snippet in hard[:25]:
            print(f"    [{label}] ...{snippet}...")
        if len(hard) > 25:
            print(f"    ... and {len(hard) - 25} more")
        failures.append(f"{len(hard)} template artifact(s)")
    else:
        print("\n  template artifacts: clean")

    soft = scan(chunks, SOFT_PATTERNS)
    if soft:
        print("\n  POSSIBLE PLACEHOLDERS (read these — some are legitimate prose)")
        for label, snippet in soft[:15]:
            print(f"    [{label}] ...{snippet}...")
        warnings.append(f"{len(soft)} possible placeholder marker(s) — confirm each "
                        "is intended copy, not leftover scaffolding")

    # Assessment reports are table-heavy and use empty paragraphs as spacers,
    # so this heuristic only means anything for the paragraph-driven plans.
    if doctype == "plan":
        empty_paras = sum(1 for p in paras if p.strip() == "")
        if empty_paras > len(paras) * 0.5 and len(paras) > 20:
            warnings.append(f"{empty_paras}/{len(paras)} paragraphs empty — "
                            "check for a section that failed to render")

    for term in expect:
        if term.lower() not in text:
            failures.append(f"expected term absent: {term!r}")
    if expect:
        found = [t for t in expect if t.lower() in text]
        print(f"  expected terms present: {len(found)}/{len(expect)}")

    # --- warnings ----------------------------------------------------------
    callout_markers = profile["callouts"]
    missing_callouts = [name for name, keys in callout_markers.items()
                        if not any(k in text for k in keys)]
    present_callouts = [n for n in callout_markers if n not in missing_callouts]
    print(f"  callouts found: {', '.join(present_callouts) or 'none'}")
    if missing_callouts:
        print(f"  callouts absent: {', '.join(missing_callouts)}")
        warnings.append("absent callouts may mean a missing client field: "
                        + ", ".join(missing_callouts))

    missing_sections = [name for name, keys in profile["sections"].items()
                        if not any(k in text for k in keys)]
    if missing_sections:
        warnings.append("sections not detected: " + ", ".join(missing_sections))

    # --- sibling comparison ------------------------------------------------
    if compare:
        _, cparas, ccells = extract(compare)
        print(f"\n  vs {Path(compare).name}: "
              f"paragraphs {len(paras)} vs {len(cparas)}, "
              f"cells {len(cells)} vs {len(ccells)}")
        if len(cparas) and len(paras) < len(cparas) * 0.5:
            warnings.append(
                f"only {len(paras)} paragraphs vs sibling's {len(cparas)} — "
                "likely a missing standard section")

    # --- report ------------------------------------------------------------
    if warnings:
        print("\n  WARNINGS (judgment required, not automatic failures)")
        for w in warnings:
            print(f"    - {w}")

    if failures:
        print("\n  FAIL: " + "; ".join(failures))
        print("  Fix by editing the build script and re-running it — "
              "never by patching the .docx.")
        return 1

    print("\n  PASS (structural)")
    return 0


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("files", nargs="+", help=".docx file(s) to audit")
    ap.add_argument("--compare", help="sibling .docx to compare section counts against")
    ap.add_argument("--type", choices=["plan", "assessment"], default=None,
                    help="document profile (auto-detected from the filename if omitted)")
    ap.add_argument("--expect", nargs="*", default=[],
                    help="terms that must appear (client name, day titles, key lifts)")
    args = ap.parse_args()

    status = 0
    for f in args.files:
        if not Path(f).exists():
            print(f"missing file: {f}")
            status = 1
            continue
        status |= audit(f, args.compare, args.expect, args.type)
    sys.exit(status)


if __name__ == "__main__":
    main()
