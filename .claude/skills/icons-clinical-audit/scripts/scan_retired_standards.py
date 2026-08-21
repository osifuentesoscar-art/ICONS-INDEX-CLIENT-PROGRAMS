#!/usr/bin/env python3
"""Scan ICONS content for clinical claims that current standards have retired.

Structural audits check whether a document rendered. This checks whether what
it says is still true. Every pattern below corresponds to a standard that was
corrected in CLAUDE.md's Research Update Log -- content written before the
correction keeps asserting it, silently, forever.

    python3 scan_retired_standards.py <paths...> [--json]

Scans .docx, .js, .json, .md. Exit 1 if any HIGH finding is present.
"""
import argparse, json, re, sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    Document = None

# (id, severity, regex, what's wrong, what current standard says)
RULES = [
    ("alst-female-optimal", "HIGH",
     r"(?:ALST|appendicular)[^.\n]{0,80}optimal|optimal[^.\n]{0,40}(?:7\.0|ALST)",
     "Claims an 'optimal' ALST tier.",
     "There is no female optimal tier. 7.0 kg/m2 is the MALE at-risk cutoff; the "
     "female cutoff is 5.5. Report ALST as a trend metric against a reference floor."),

    ("alst-three-tier", "HIGH",
     r"5\.5[^.\n]{0,40}(?:6\.99|7\.0)[^.\n]{0,40}optimal|normal[^.\n]{0,20}monitor",
     "Uses the retired three-tier women's ALST ladder.",
     "Two tiers only for women: <5.5 At-Risk, >=5.5 within normal reference range."),

    ("vfa-risk-band", "HIGH",
     r"(?:VFA|visceral)[^.\n]{0,60}(?:very low|low|moderate|high)\s+risk|"
     r"(?:very low|low|moderate|high)\s+risk[^.\n]{0,40}(?:cm2|cm²)",
     "Assigns a VFA risk band.",
     "The four-tier VFA table is retired -- no consensus body endorses a single "
     "cutoff, and Styku validated VFA against DXA in kg, not CT in cm2. Track as "
     "a personal trend; use waist circumference as the clinical-facing metric."),

    ("asymmetry-absolute", "MEDIUM",
     r"0\.5\s*(?:lb|lbs|pound)[^.\n]{0,40}(?:trigger|threshold|gap|asymmetr)|"
     r"asymmetr[^.\n]{0,40}0\.5\s*lb",
     "Uses the retired 0.5 lb absolute asymmetry trigger.",
     "Trigger is a RELATIVE gap >=10%. The device's own error on leg lean mass is "
     "2-3x larger than 0.5 lb, so the old trigger fired on noise."),

    ("protein-age-tier", "MEDIUM",
     r"(?:40\+|50\+)\s*tier|(?:1\.8\s*[-–]\s*2\.0|2\.0\s*[-–]\s*2\.2)\s*g/kg",
     "Uses an age-banded protein escalation.",
     "Protein is context-keyed, not age-keyed: 1.6 g/kg baseline, up to 2.2 for "
     "energy deficit, heavy training load, or ALST At-Risk."),

    ("rescan-single-clock", "MEDIUM",
     r"(?:reassess|re-?assess|re-?test)[^.\n]{0,50}(?:every\s*)?8\s*(?:-|–|to)?\s*(?:12)?\s*week",
     "States a single 8-week reassessment clock.",
     "Two clocks: strength re-checked every 4 weeks; Styku rescan on its own "
     "8-12 week cycle. Do not collapse them."),

    ("failure-training", "MEDIUM",
     r"train(?:ing)?\s+to\s+failure\s+(?:is|for)\s+(?:superior|better|optimal)|to\s+failure\s+every\s+set",
     "Recommends training to failure as superior.",
     "ACSM 2026: failure training does not consistently outperform RIR-based "
     "loading. Default 2 RIR on primaries, 1 RIR on hypertrophy accessories."),

    ("acl-screen-gated", "MEDIUM",
     r"(?:if|when)[^.\n]{0,40}(?:valgus|knee collapse)[^.\n]{0,40}(?:screen|observed|positive)[^.\n]{0,60}(?:then|add|prescribe)",
     "Gates the neuromuscular circuit behind a positive valgus screen.",
     "Visual screening cannot predict future ACL injury. The circuit is universal "
     "and dosed (20-30 min, 1-2x/week, sustained past 6 months)."),

    ("pelvic-no-breath-hold", "MEDIUM",
     r"never\s+hold\s+your\s+breath|no\s+breath.?holding|avoid\s+breath.?holding",
     "States a blanket no-breath-holding rule.",
     "POGP 2024: a permanent prohibition is not appropriate. Pressure strategy "
     "scales with load; a brief controlled brace on the heaviest sets is normal. "
     "What stays wrong is SUSTAINED holding across reps or a whole set."),

    ("liftmor-80", "LOW",
     r"(?:>=|≥|at least\s*)80\s*%\s*1RM[^.\n]{0,60}(?:bone|LIFTMOR|5x5)",
     "States LIFTMOR intensity as >=80% 1RM.",
     "The trial prescribed 5x5 at >85% 1RM, supervised, with a risk-stratification "
     "gate and a technique-first ramp-in."),

    ("copenhagen-retracted", "HIGH",
     r"Scandinavian Journal of Medicine[^.\n]{0,60}Copenhagen|Copenhagen[^.\n]{0,80}2025[^.\n]{0,40}meta-analys",
     "Possible citation of the RETRACTED 2025 Copenhagen meta-analysis.",
     "That paper was formally retracted (notice 2026;36(4):e70287). Cite Haroy et "
     "al., BJSM 2019;53:150 instead."),
]

NEGATORS = (r"(not|never|no longer|retired|superseded|don'?t|avoid|instead of|"
            r"rather than|incorrect|wrong|corrected|used to|previously|"
            r"deprecated|anything that|must not|would be)")


def read(path):
    p = Path(path)
    if p.suffix == ".docx":
        if Document is None:
            return ""
        d = Document(p)
        parts = [x.text for x in d.paragraphs]
        for t in d.tables:
            for r in t.rows:
                for c in r.cells:
                    parts.append(c.text)
        return "\n".join(parts)
    try:
        return p.read_text(errors="ignore")
    except Exception:
        return ""


def scan_text(text, path):
    findings = []
    for rid, sev, pat, problem, correct in RULES:
        for m in re.finditer(pat, text, re.I):
            # Content that names a retired standard in order to REJECT it is
            # correct content, and flagging it trains the reader to ignore this
            # tool. But the negation check has to be tight: a wide window
            # almost always contains some "not", which suppresses real
            # findings. Scope it to the sentence, and require the negator to
            # come BEFORE the matched phrase -- "not At-Risk" appearing after
            # "optimal threshold" does not negate the optimal claim.
            sent_start = max((text.rfind(c, 0, m.start()) for c in ".!?\n"), default=-1) + 1
            nxt = [i for i in (text.find(c, m.end()) for c in ".!?\n") if i != -1]
            sent_end = min(nxt) if nxt else len(text)
            preceding = text[sent_start:m.start()]
            # A short forward window too: reference material often names a
            # standard and then says it is retired ("...2.0-2.2 at 50+) is
            # retired"). Keep it narrow -- a wide forward window swallows the
            # real findings again.
            following = text[m.end():min(sent_end, m.end() + 80)]
            if re.search(NEGATORS, preceding, re.I) or re.search(NEGATORS, following, re.I):
                continue
            window = text[sent_start:sent_end]
            line = text[:m.start()].count("\n") + 1
            findings.append({
                "file": str(path), "line": line, "rule": rid, "severity": sev,
                "match": m.group()[:90].replace("\n", " "),
                "context": window[:260].replace("\n", " "),
                "problem": problem, "current_standard": correct,
            })
            break  # one finding per rule per file is enough to act on
    return findings


def main():
    ap = argparse.ArgumentParser(description=__doc__,
            formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("paths", nargs="+")
    ap.add_argument("--json", action="store_true")
    args = ap.parse_args()

    targets = []
    for raw in args.paths:
        p = Path(raw)
        if p.is_dir():
            for ext in ("*.docx", "*.js", "*.json", "*.md"):
                targets += [f for f in p.rglob(ext) if "node_modules" not in str(f)]
        elif p.exists():
            targets.append(p)

    findings = []
    for t in targets:
        findings += scan_text(read(t), t)

    if args.json:
        print(json.dumps(findings, indent=2))
        return 1 if any(f["severity"] == "HIGH" for f in findings) else 0

    if not findings:
        print(f"Scanned {len(targets)} file(s). No retired-standard claims found.")
        return 0

    order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    findings.sort(key=lambda f: (order[f["severity"]], f["file"]))
    print(f"Scanned {len(targets)} file(s). {len(findings)} finding(s).\n")
    for f in findings:
        print(f"[{f['severity']}] {f['rule']}  —  {f['file']}:{f['line']}")
        print(f"    matched : {f['match']}")
        print(f"    problem : {f['problem']}")
        print(f"    current : {f['current_standard']}")
        print()
    highs = sum(1 for f in findings if f["severity"] == "HIGH")
    if highs:
        print(f"{highs} HIGH finding(s) — these are clinical claims that are now wrong.")
    return 1 if highs else 0


if __name__ == "__main__":
    sys.exit(main())
